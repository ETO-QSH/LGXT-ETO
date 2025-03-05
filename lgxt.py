import os
import requests
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from reportlab.lib import colors
from reportlab.lib.units import inch

base_url = "http://lgxt.wutp.com.cn/api"
headers = {
    'Accept': '*/*',
    'Accept-Language': 'zh-CN,zh;q=0.9',
    'Content-Type': 'application/x-www-form-urlencoded',
}

session = requests.Session()

def login(username, password):
    login_url = f"{base_url}/login"
    data = {'loginName': username, 'password': password}
    try:
        response = session.post(login_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            token = result['data']
            headers['Authorization'] = token
            session.headers.update({'Authorization': token})
            return True, "欢迎回来！"
        else:
            return False, f"登录失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"

def get_my_courses():
    my_courses_url = f"{base_url}/myCourses"
    try:
        response = session.post(my_courses_url, headers=headers, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            courses = result['data']
            return True, courses
        else:
            return False, f"获取课程列表失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"

def get_course_works(course_id):
    my_course_works_url = f"{base_url}/myCourseWorks"
    data = {'courseId': course_id}
    try:
        response = session.post(my_course_works_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            works = result['data']
            return True, works
        else:
            return False, f"获取课程作业失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"

def get_questions(work_id):
    show_questions_url = f"{base_url}/showQuestions"
    data = {'workId': work_id}
    try:
        response = session.post(show_questions_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            questions = result['data']
            return True, questions
        else:
            return False, f"获取题目失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def submit_answer(work_id, grade):
    submit_answer_url = f"{base_url}/submitAnswer"
    data = {'grade': grade, 'workId': work_id}
    try:
        response = session.post(submit_answer_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            return True, f"答案提交成功，成绩：{grade}"
        else:
            return False, f"答案提交失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def collect_all_questions(work_id, max_iterations=100):
    collected_questions = {}
    no_new_questions_count = 0
    for i in range(max_iterations):
        success, questions = get_questions(work_id)
        if success:
            new_question_found = False
            for question in questions:
                question_id = question.get('id')
                if question_id not in collected_questions:
                    collected_questions[question_id] = question
                    new_question_found = True
            if new_question_found:
                no_new_questions_count = 0
            else:
                no_new_questions_count += 1
                if no_new_questions_count >= 10:
                    break
        else:
            break
    return collected_questions

def save_questions_to_word(collected_questions, assignment_folder, work_name, export_answers=True):
    document = Document()

    style = document.styles['Normal']
    font_style = style.font
    font_style.name = '宋体'
    font_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    document.add_heading(work_name, 0)

    images_folder = os.path.join(assignment_folder, '题目图片')

    sorted_question_ids = sorted(collected_questions.keys(), key=lambda x: int(x))

    for idx, question_id in enumerate(sorted_question_ids):
        question = collected_questions[question_id]
        name = question.get('name', 'N/A')
        answer = question.get('answer', 'N/A')

        document.add_heading(f'题目 {idx + 1}: {name}', level=2)

        image_path = os.path.join(images_folder, f"{question_id}.png")
        if os.path.exists(image_path):
            document.add_picture(image_path, width=Inches(5))
        else:
            document.add_paragraph("（无图片）")

        if export_answers:
            answer_paragraph = document.add_paragraph("答案：")
            answer_run = answer_paragraph.add_run(answer)
            answer_run.font.color.rgb = RGBColor(255, 0, 0)

    docx_path = os.path.join(f'{work_name}.docx')
    document.save(docx_path)

def save_questions_to_pdf(collected_questions, assignment_folder, work_name, export_answers=True):
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Image as RLImage, Spacer
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader

    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ChineseTitle',
        parent=styles['Title'],
        fontName='STSong-Light',
        fontSize=20,
        leading=24,
        alignment=1,
        textColor=colors.HexColor('#333333')
    ))
    styles.add(ParagraphStyle(
        name='ChineseHeading1',
        parent=styles['Heading1'],
        fontName='STSong-Light',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#555555')
    ))
    styles.add(ParagraphStyle(
        name='Chinese',
        fontName='STSong-Light',
        fontSize=12,
        leading=18,
        textColor=colors.HexColor('#000000')
    ))
    normal_style = styles['Chinese']

    def add_page_number(canvas, _doc):
        page_num = canvas.getPageNumber()
        text = f"第 {page_num} 页"
        canvas.setFont('STSong-Light', 9)
        canvas.drawRightString(A4[0] - 50, 15, text)

    left_margin = 40
    right_margin = 40
    top_margin = 60
    bottom_margin = 60

    doc = BaseDocTemplate(
        f'{work_name}.pdf',
        pagesize=A4,
        rightMargin=right_margin,
        leftMargin=left_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin
    )

    page_width, page_height = A4
    frame_width = page_width - left_margin - right_margin
    frame_height = page_height - top_margin - bottom_margin

    frame = Frame(left_margin, bottom_margin, frame_width, frame_height, id='normal')

    template = PageTemplate(id='test', frames=frame, onPage=add_page_number)
    doc.addPageTemplates([template])

    elements = [Paragraph(f'{work_name}', styles['ChineseTitle']), Spacer(1, 0.3 * inch)]

    images_folder = os.path.join(assignment_folder, '题目图片')

    sorted_question_ids = sorted(collected_questions.keys(), key=lambda x: int(x))

    for idx, question_id in enumerate(sorted_question_ids):
        question = collected_questions[question_id]
        name = question.get('name', 'N/A')
        answer = question.get('answer', 'N/A')

        title_text = f'题目 {idx + 1}: {name}'
        elements.append(Paragraph(title_text, styles['ChineseHeading1']))
        elements.append(Spacer(1, 0.1 * inch))

        image_path = os.path.join(images_folder, f"{question_id}.png")
        if os.path.exists(image_path):
            img = ImageReader(image_path)
            img_width, img_height = img.getSize()
            aspect = img_height / float(img_width)

            img_display_width = doc.width * 0.8
            img_display_height = img_display_width * aspect

            elements.append(RLImage(image_path, width=img_display_width, height=img_display_height))
        else:
            elements.append(Paragraph("（无图片）", normal_style))

        if export_answers:
            answer_paragraph = Paragraph(f"答案：<font color='red'>{answer}</font>", normal_style)
            elements.append(answer_paragraph)

        elements.append(Spacer(1, 0.2 * inch))

    doc.build(elements)

def save_question(question, assignment_folder):
       question_id = question.get('id', 'N/A')
       imgurl = question.get('imgurl', 'N/A')

       images_folder = os.path.join(assignment_folder, '题目图片')
       if not os.path.exists(images_folder):
           os.makedirs(images_folder)

       if imgurl and imgurl != 'N/A':
           try:
               response = session.get(imgurl)
               response.raise_for_status()
               image_data = response.content
               image_path = os.path.join(images_folder, f"{question_id}.png")
               with open(image_path, 'wb') as img_file:
                   img_file.write(image_data)
           except Exception as e:
               print(f"无法下载题目 {question_id} 的图片：{e}")
